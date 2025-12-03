# ============================================================
# SmartTest Architect — FINAL MERGED app.py (with DOCX MST export)
# Added: create_mst_docx(...) to produce DOCX identical-in-structure to the PDF
# ============================================================

import streamlit as st
from groq import Groq
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import re, random
from difflib import SequenceMatcher
import os

# ---------------- CONFIG ----------------
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = "llama-3.1-8b-instant"
client = None
try:
    client = Groq(api_key=GROQ_API_KEY)
except Exception:
    client = None

# ---------------- UNICODE FIX ----------------
def safe_unicode(text):
    if text is None:
        return ""
    t = str(text)
    t = t.replace("\u2014", "-").replace("\u2013", "-")
    t = t.replace("\u2018", "'").replace("\u2019", "'")
    t = t.replace("\u201c", '"').replace("\u201d", '"')
    t = t.replace("\u2026", "...").replace("\u00A0", " ")
    t = re.sub(r"[\x00-\x1F\x7F]", "", t)
    return t.encode("latin-1", "replace").decode("latin-1")

# ---------------- SAFE GROQ CALL ----------------
def groq_safe(msgs, max_tokens=600, temperature=0.15):
    if client is None:
        return "GROQ_NOT_CONFIGURED"
    try:
        res = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=msgs,
            temperature=temperature,
            max_tokens=max_tokens
        )
        return res.choices[0].message.content
    except Exception as e:
        return f"GROQ_ERROR: {str(e)}"

# ---------------- UNIT PARSER ----------------
def parse_units_from_syllabus(text):
    text = text.replace("–", "-").replace("—", "-")
    text = text.replace("\r", "").strip()

    pattern = re.compile(
        r"((UNIT|MODULE|CHAPTER)\s*[-:. )]*\s*(I|II|III|IV|V|VI|VII|VIII|IX|X|\d+).*?)(?=(?:\n\s*(UNIT|MODULE|CHAPTER)\b)|\Z)",
        re.IGNORECASE | re.DOTALL
    )

    matches = re.findall(pattern, text)
    units = []

    if matches:
        counter = 1
        for block in matches:
            full_block = block[0]
            cleaned = re.sub(r"(UNIT|MODULE|CHAPTER)\s*[-:. )]*\s*\w+", "", full_block, flags=re.I).strip()
            units.append({"title": f"UNIT - {counter}", "content": cleaned})
            counter += 1
    else:
        parts = [p.strip() for p in text.split("\n") if p.strip()]
        for i, p in enumerate(parts):
            units.append({"title": f"UNIT - {i+1}", "content": p})

    return units

# ---------------- CO PARSER ----------------
def parse_co_list(text):
    lines = text.splitlines()
    co_list = []
    current = ""

    for ln in lines:
        ln = ln.strip()
        if not ln:
            continue

        if re.match(r"^\s*\d+\s*\.", ln):
            if current.strip():
                co_list.append(current.strip())
            current = re.sub(r"^\s*\d+\s*\.\s*", "", ln).strip()
        else:
            current += " " + ln

    if current.strip():
        co_list.append(current.strip())

    return co_list

# ---------------- DUPLICATE PREVENTION ----------------
if "USED_Q" not in st.session_state:
    st.session_state["USED_Q"] = set()
if "USED_TOPICS" not in st.session_state:
    st.session_state["USED_TOPICS"] = set()

def norm(txt):
    t = re.sub(r"[^a-z0-9 ]", "", str(txt).lower())
    t = re.sub(r"\s+", " ", t).strip()
    return t

def is_dup(q):
    key = norm(q)
    if key in st.session_state["USED_Q"]:
        return True
    for old in st.session_state["USED_Q"]:
        if SequenceMatcher(None, old, key).ratio() > 0.80:
            return True
    return False

def add_dup(q):
    st.session_state["USED_Q"].add(norm(q))

# ---------------- RBT / KC ----------------
RBT_KEYWORDS = {
    "L1": ["define", "what is", "list", "state", "identify"],
    "L2": ["explain", "describe", "discuss", "summarize"],
    "L3": ["apply", "solve", "use", "demonstrate"],
    "L4": ["analyze", "compare", "differentiate"],
    "L5": ["evaluate", "justify", "argue"],
    "L6": ["create", "design", "develop"]
}

KC_KEYWORDS = {
    "F": ["define", "list", "state", "identify"],
    "C": ["explain", "describe", "discuss", "compare"],
    "P": ["apply", "solve", "use", "implement"],
    "M": ["justify", "reflect", "evaluate"]
}

def heuristic_rbt_kc(q):
    ql = q.lower()
    for lvl, words in RBT_KEYWORDS.items():
        for w in words:
            if w in ql:
                for kc, kws in KC_KEYWORDS.items():
                    if any(kw in ql for kw in kws):
                        return lvl, kc
                return lvl, "C"
    return "L3", "C"

# ---------------- CO DETECTOR ----------------
def detect_co(question, co_lines):
    if not co_lines:
        rbt, kc = heuristic_rbt_kc(question)
        return "CO1", "Default", f"CO1, {rbt}, {kc}"

    best_idx = 0
    best_score = -1
    qnorm = re.sub(r'\W+', ' ', question.lower())

    for i, co in enumerate(co_lines):
        cnorm = re.sub(r'\W+', ' ', co.lower())
        score = SequenceMatcher(None, qnorm, cnorm).ratio()
        if score > best_score:
            best_score = score
            best_idx = i

    label = f"CO{best_idx+1}"
    co_text = co_lines[best_idx]
    rbt, kc = heuristic_rbt_kc(question)
    return label, co_text, f"{label}, {rbt}, {kc}"

# ---------------- TOPIC EXTRACTOR ----------------
def extract_topics(unit_content):
    text = unit_content.replace("\n", " ").strip()
    text = re.sub(r"\s+", " ", text)

    text = re.sub(r"\b\d+\s*(hrs|hours|hr)\b", "", text, flags=re.I)
    text = re.sub(r"\b\d+\b", "", text)

    parts = re.split(r"[;,:\.\(\)\-–—/]", text)

    stop = {"and", "or", "etc", "the", "a", "an", "to"}
    topics = []

    for p in parts:
        p = p.strip()
        if len(p) < 3:
            continue
        if p.lower() in stop:
            continue
        topics.append(p)

    topics = list(dict.fromkeys(topics))

    if not topics:
        topics = ["General Concept"]

    return topics

# ---------------- QUESTION GENERATOR ----------------
# Now supports question_type: Generic / Semantic / Logical / Syntax / Theory

def make_question(topic, kind, difficulty, question_type="Generic", unit_context=""):

    # Basic template pools (kept for fallback and faster local generation)
    EASY_TEMPLATES = [
        "Define {topic} with a suitable example.",
        "What is {topic}? Explain in simple terms.",
        "List the key features of {topic}.",
        "State the meaning of {topic} and mention its purpose."
    ]

    MEDIUM_TEMPLATES = [
        "Explain {topic} in detail with relevant points.",
        "Describe the working of {topic} with an illustration.",
        "Discuss the importance of {topic} in the subject domain.",
        "Differentiate {topic} from closely related concepts."
    ]

    HARD_TEMPLATES = [
        "Analyze {topic} and highlight its practical applications.",
        "Evaluate the role of {topic} in modern systems with justification.",
        "Design a scenario where {topic} can be effectively applied.",
        "Critically examine {topic} and compare it with alternatives."
    ]

    # If user asked for specialized question types, try to use the AI (GROQ) for high-quality generation
    if question_type != "Generic" and client is not None:
        try:
            prompt = prompt = (
                f"You are a professional university exam-setter. "
                f"Create ONE {question_type} question STRICTLY from the following UNIT SYLLABUS context:\n\n"
                f"{unit_context}\n\n"
                f"The question MUST:\n"
                f"Rules:\n"
                f"- Be 100% from the taught syllabus (NO out-of-syllabus content). Do NOT start the question with words like “Consider”, “Given”, “Suppose”, “Assume”.Start directly with the topic or verb (Explain/Analyze/Define/Discuss).\n"
                f"- Be relevant to the topic: '{topic}'\n"
                f"- Match an Indian university MST style\n"
                f"- Sound like a real teacher who taught this unit has created it\n"
                f"- Maintain academic tone\n"
                f"- Keep language academic and concise\n"
                f"- Be ONLY 1–2 lines long (concise)\n"
                f"- Avoid extra explanation\n\n"
                f"Difficulty: {difficulty}\n"
                f"Format: {kind}\n\n"
                f"Give ONLY the question text."
                )


            out = groq_safe([
                {"role": "system", "content": "You are a strict academic question generator."},
                {"role": "user", "content": prompt}
            ], max_tokens=200, temperature=0.2)

            if out and not out.startswith("GROQ_ERROR") and out != "GROQ_NOT_CONFIGURED":
                qtext = safe_unicode(out).strip().replace('\n', ' ')
                # Basic normalization
                if not qtext.endswith('.') and not qtext.endswith('?'):
                    qtext = qtext + '.'
                # Add kind-specific suffixes
                if kind == "medium":
                    qtext += " Explain each point clearly."
                elif kind == "long":
                    qtext += " Support your answer with examples or diagrams."
                return qtext
        except Exception:
            pass

    # fallback: local templates
    if difficulty == "Easy":
        template = random.choice(EASY_TEMPLATES)
    elif difficulty == "Medium":
        template = random.choice(MEDIUM_TEMPLATES)
    else:
        template = random.choice(HARD_TEMPLATES)

    if kind == "medium":
        template += " Explain each point clearly."
    elif kind == "long":
        template += " Support your answer with examples or diagrams."

    # For certain question types, tweak phrasing locally
    if question_type == "Logical":
        template = "Given the following scenario related to {topic}, analyze and solve the problem. Provide steps and justification.".format(topic=topic)
    elif question_type == "Syntax":
        template = "Write correct syntax or code snippet for {topic} and explain why it works.".format(topic=topic)
    elif question_type == "Theory":
        template = "Discuss the theory behind {topic} and its historical or conceptual development.".format(topic=topic)
    elif question_type == "Semantic":
        # If semantic requested but no AI available, use detailed conceptual template
        template = "Explain the semantic meaning of {topic} in context and provide examples where it applies.".format(topic=topic)

    return template.format(topic=topic)


def random_rbt_level(section_list_length):
    # possible RBT levels
    levels = ["L1", "L2", "L3", "L4", "L5", "L6"]
    # shuffle to get randomness
    random.shuffle(levels)
    # return levels based on how many questions in that section
    return levels[:section_list_length]

def generate_questions(unit_title, unit_content, count, kind="short", difficulty="Medium", co_lines=None, question_type="Generic"):
    topics = extract_topics(unit_content)
    random.shuffle(topics)

    final_topics = []

    for t in topics:
        if t.lower() not in st.session_state["USED_TOPICS"]:
            final_topics.append(t)
            st.session_state["USED_TOPICS"].add(t.lower())
        if len(final_topics) == count:
            break

    while len(final_topics) < count:
        extra = f"{unit_title} concept {len(final_topics)+1}"
        if extra.lower() not in st.session_state["USED_TOPICS"]:
            st.session_state["USED_TOPICS"].add(extra.lower())
            final_topics.append(extra)

    questions = []
    for t in final_topics:
        q = make_question(t, kind, difficulty, question_type=question_type, unit_context=unit_content)
        q = re.sub(r'(\b[A-Za-z0-9_]+\b)\s*\?\s*(\b[A-Za-z0-9_]+\b)', r'\1 → \2', q)
        q = re.sub(r'(\b[A-Za-z0-9_]+)\?([A-Za-z0-9_]+\b)', r'\1 → \2', q)
        q = re.sub(r'(\b[A-Za-z0-9_]+)\?\s+([A-Za-z0-9_]+\b)', r'\1 → \2', q)

        label, co_text, tag = detect_co(q, co_lines or [])

        if is_dup(q):
            q += " (explain briefly)"

        add_dup(q)
        questions.append({"q": q, "tag": tag})

    return questions

# ---------------- ANSWER GENERATOR ----------------
def generate_answer(question_text, kind="short"):
    if client is not None:
        try:
            out = groq_safe([
                {"role": "system", "content": "You write concise academic model answers."},
                {"role": "user", "content": f"Write a {kind} model answer for: {question_text}"}
            ], max_tokens=300)
            if out and not out.startswith("GROQ_ERROR") and out != "GROQ_NOT_CONFIGURED":
                return safe_unicode(out)
        except Exception:
            pass
    return f"Model answer ({kind}): Brief explanation of {question_text}"

# ============================================================
# PDF / DOCX BUILDERS
# ============================================================
class PDF(FPDF):
    def header(self):
        pass
    def footer(self):
        self.set_y(-12)
        self.set_font("Times", "I", 8)
        self.cell(0, 6, f"Page {self.page_no()}", align="R")
    def safe(self, txt):
        return safe_unicode(txt)

# --- existing PDF builder kept as-is ---
def create_mst_pdf(sections, header):
    pdf = PDF("P", "mm", "A4")
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()

    # HEADER
    pdf.set_font("Times", "B", 12)
    pdf.cell(0, 6, pdf.safe(header.get("university", "")), ln=True, align="C")
    pdf.set_font("Times", "B", 11)
    pdf.cell(0, 6, pdf.safe(header.get("college", "")), ln=True, align="C")
    pdf.set_font("Times", "", 10)
    pdf.cell(0, 5, pdf.safe(header.get("department", "")), ln=True, align="C")
    pdf.ln(3)

    # MAIN INFO TABLE
    pdf.set_font("Times", "", 9)
    W1, W2, W3, W4 = 35, 55, 35, 55
    H = 6
    rows = [
        ("Academic Year", header.get("academic_year", ""), "Session", header.get("session", "")),
        ("Program", header.get("program", ""), "Semester", header.get("semester", "")),
        ("Course Code", header.get("course_code", ""), "Course Title", header.get("course_title", "")),
        ("Mid Semester Test (MST)", header.get("mid_mst", ""), "Course Coordinator", header.get("coordinator", "")),
        ("Max. Marks", header.get("max_marks", ""), "Time Duration", header.get("time_duration", "")),
        ("Date of MST", header.get("date_of_MST", ""), "Roll Number", header.get("roll_number", "")),
    ]
    for r in rows:
        pdf.cell(W1, H, pdf.safe(r[0]), 1, 0)
        pdf.cell(W2, H, pdf.safe(r[1]), 1, 0)
        pdf.cell(W3, H, pdf.safe(r[2]), 1, 0)
        pdf.cell(W4, H, pdf.safe(r[3]), 1, 1)

    pdf.set_font("Times", "B", 9)
    pdf.cell(0, 6, "Note: Attempt all questions.", ln=True)
    pdf.ln(3)

    # QUESTION TABLE HEADER
    pdf.set_font("Times", "B", 9)
    QW, QSW, COW, MW = 12, 120, 38, 12
    pdf.cell(QW, 7, "Q.No", 1, 0, "C")
    pdf.cell(QSW, 7, "Question", 1, 0, "C")
    pdf.cell(COW, 7, "CO/RBT/KC", 1, 0, "C")
    pdf.cell(MW, 7, "Marks", 1, 1, "C")

    def draw_row(qno, question, tag, marks):
        pdf.set_font("Times", "", 9)
        lh = 4
        q_lines = pdf.multi_cell(QSW, lh, pdf.safe(question), split_only=True)
        t_lines = pdf.multi_cell(COW, lh, pdf.safe(tag), split_only=True)
        m_lines = pdf.multi_cell(MW, lh, pdf.safe(str(marks)), split_only=True)
        n = max(len(q_lines), len(t_lines), len(m_lines))
        h = max(8, n * lh)

        if pdf.get_y() + h + 20 > pdf.h - pdf.b_margin:
            pdf.add_page()
            pdf.set_font("Times", "B", 9)
            pdf.cell(QW, 7, "Q.No", 1, 0, "C")
            pdf.cell(QSW, 7, "Question", 1, 0, "C")
            pdf.cell(COW, 7, "CO/RBT/KC", 1, 0, "C")
            pdf.cell(MW, 7, "Marks", 1, 1, "C")

        x = pdf.get_x()
        y = pdf.get_y()

        pdf.rect(x, y, QW, h)
        pdf.rect(x + QW, y, QSW, h)
        pdf.rect(x + QW + QSW, y, COW, h)
        pdf.rect(x + QW + QSW + COW, y, MW, h)

        pdf.set_xy(x, y)
        pdf.multi_cell(QW, h / (n if n else 1), pdf.safe(f"Q{qno}"), align="C")

        pdf.set_xy(x + QW, y)
        pdf.multi_cell(QSW, lh, pdf.safe(question))

        pdf.set_xy(x + QW + QSW, y)
        pdf.multi_cell(COW, lh, pdf.safe(tag), align="C")

        pdf.set_xy(x + QW + QSW + COW, y)
        pdf.multi_cell(MW, lh, pdf.safe(str(marks)), align="C")

        pdf.set_xy(x, y + h)

    # SECTION A
    pdf.set_font("Times", "B", 9)
    pdf.cell(QW + QSW + COW + MW, 7, "Section-A (4 × 2)", 1, 1, "C")
    pdf.set_font("Times", "", 9)
    qno = 1
    for q in sections.get("A", []):
        draw_row(qno, q.get("q", ""), q.get("tag", ""), q.get("marks", 2))
        qno += 1

    # ---------------- RANDOM RBT MIX FOR SECTION A ----------------
    A_levels = random_rbt_level(len(sections["A"]))
    for i, q in enumerate(sections["A"]):
        old_tag = q["tag"].split(",")
        kc = old_tag[-1].strip()
        co = old_tag[0].strip()
        rbt = A_levels[i % len(A_levels)]
        q["tag"] = f"{co}, {rbt}, {kc}"

    # SECTION B
    pdf.set_font("Times", "B", 9)
    pdf.cell(QW + QSW + COW + MW, 7, "Section-B (2 × 4)", 1, 1, "C")
    pdf.set_font("Times", "", 9)
    for q in sections.get("B", []):
        draw_row(qno, q.get("q", ""), q.get("tag", ""), q.get("marks", 4))
        qno += 1

    # ---------------- RANDOM RBT MIX FOR SECTION B ----------------
    B_levels = random_rbt_level(len(sections["B"]))
    for i, q in enumerate(sections["B"]):
        old_tag = q["tag"].split(",")
        kc = old_tag[-1].strip()
        co = old_tag[0].strip()
        rbt = B_levels[i % len(B_levels)]
        q["tag"] = f"{co}, {rbt}, {kc}"

    # SECTION C
    pdf.set_font("Times", "B", 9)
    pdf.cell(QW + QSW + COW + MW, 7, "Section-C (1 × 8)", 1, 1, "C")
    pdf.set_font("Times", "", 9)
    for q in sections.get("C", []):
        draw_row(qno, q.get("q", ""), q.get("tag", ""), q.get("marks", 8))
        qno += 1

    # ---------------- RANDOM RBT MIX FOR SECTION C ----------------
    C_levels = random_rbt_level(len(sections["C"]))
    for i, q in enumerate(sections["C"]):
        old_tag = q["tag"].split(",")
        kc = old_tag[-1].strip()
        co = old_tag[0].strip()
        rbt = C_levels[i % len(C_levels)]
        q["tag"] = f"{co}, {rbt}, {kc}"


    pdf.ln(8)

    # COURSE OUTCOMES TABLE
    pdf.set_font("Times", "B", 10)
    pdf.cell(0, 6, "COURSE OUTCOMES (CO): Students will be able to", ln=True)
    pdf.set_font("Times", "", 9)
    WCO1, WCO2 = 10, 178
    for i, co in enumerate(header.get("co_lines", []), start=1):
        pdf.cell(WCO1, 7, str(i), 1, 0, "C")
        pdf.cell(WCO2, 7, pdf.safe(co), 1, 1)

    pdf.ln(8)

    # RBT TABLE
    pdf.set_font("Times", "B", 10)
    pdf.cell(0, 6, "RBT Classification", ln=True)
    pdf.set_font("Times", "", 9)
    pdf.cell(40, 7, "Level", 1, 0, "C")
    for lvl in ["L1", "L2", "L3", "L4", "L5", "L6"]:
        pdf.cell(25, 7, lvl, 1, 0, "C")
    pdf.ln()
    pdf.cell(40, 7, "Description", 1, 0, "C")
    for nm in ["Remembering", "Understanding", "Applying", "Analyzing", "Evaluating", "Creating"]:
        pdf.cell(25, 7, nm, 1, 0, "C")
    pdf.ln(10)

    # KC TABLE
    pdf.set_font("Times", "B", 10)
    pdf.cell(0, 6, "Knowledge Category", ln=True)
    pdf.set_font("Times", "", 9)
    pdf.cell(40, 7, "KC Code", 1, 0, "C")
    for kc in ["F", "C", "P", "M"]:
        pdf.cell(30, 7, kc, 1, 0, "C")
    pdf.ln()
    pdf.cell(40, 7, "KC Name", 1, 0, "C")
    for nm in ["Factual", "Conceptual", "Procedural", "Meta-Cognitive"]:
        pdf.cell(30, 7, nm, 1, 0, "C")

    return pdf.output(dest="S").encode("latin-1", "replace")

# --- NEW: DOCX MST builder (structure mirrors the PDF layout) ---
def create_mst_docx(sections, header):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # ---------------- HEADER ----------------
    p = doc.add_paragraph()
    r = p.add_run(safe_unicode(header.get('university', '')))
    r.bold = True; r.font.size = Pt(14)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    r = p.add_run(safe_unicode(header.get('college', '')))
    r.bold = True; r.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph(safe_unicode(header.get('department', '')))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    # ---------------- MAIN INFO TABLE ----------------
    rows = [
        ("Academic Year", header.get("academic_year", ""), "Session", header.get("session", "")),
        ("Program", header.get("program", ""), "Semester", header.get("semester", "")),
        ("Course Code", header.get("course_code", ""), "Course Title", header.get("course_title", "")),
        ("Mid Semester Test (MST)", header.get("mid_mst", ""), "Course Coordinator", header.get("coordinator", "")),
        ("Max. Marks", header.get("max_marks", ""), "Time Duration", header.get("time_duration", "")),
        ("Date of MST", header.get("date_of_MST", ""), "Roll Number", header.get("roll_number", "")),
    ]

    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"

    for i, r in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = safe_unicode(r[0])
        cells[1].text = safe_unicode(r[1])
        cells[2].text = safe_unicode(r[2])
        cells[3].text = safe_unicode(r[3])

    doc.add_paragraph("Note: Attempt all questions.")
    doc.add_paragraph("")

    # ---------------- QUESTION TABLE ----------------
    qtable = doc.add_table(rows=1, cols=4)
    qtable.style = "Table Grid"

    hdr = qtable.rows[0].cells
    hdr[0].text = "Q.No"
    hdr[1].text = "Question"
    hdr[2].text = "CO/RBT/KC"
    hdr[3].text = "Marks"

    qno = 1

    # SECTION A
    row = qtable.add_row().cells
    row[1].text = "Section-A (4 × 2)"

    for q in sections.get("A", []):
        r = qtable.add_row().cells
        r[0].text = f"Q{qno}"
        r[1].text = safe_unicode(q["q"])
        r[2].text = safe_unicode(q["tag"])
        r[3].text = str(q["marks"])
        qno += 1

    # SECTION B
    row = qtable.add_row().cells
    row[1].text = "Section-B (2 × 4)"

    for q in sections.get("B", []):
        r = qtable.add_row().cells
        r[0].text = f"Q{qno}"
        r[1].text = safe_unicode(q["q"])
        r[2].text = safe_unicode(q["tag"])
        r[3].text = str(q["marks"])
        qno += 1

    # SECTION C
    row = qtable.add_row().cells
    row[1].text = "Section-C (1 × 8)"

    for q in sections.get("C", []):
        r = qtable.add_row().cells
        r[0].text = f"Q{qno}"
        r[1].text = safe_unicode(q["q"])
        r[2].text = safe_unicode(q["tag"])
        r[3].text = str(q["marks"])
        qno += 1

    doc.add_paragraph("")

    # ---------------- CO TABLE ----------------
    doc.add_paragraph("COURSE OUTCOMES (CO): Students will be able to")
    if header.get("co_lines"):
        cot = doc.add_table(rows=1, cols=2)
        cot.style = "Table Grid"
        cot.rows[0].cells[0].text = "No."
        cot.rows[0].cells[1].text = "Outcome"

        for i, co in enumerate(header["co_lines"], start=1):
            r = cot.add_row().cells
            r[0].text = str(i)
            r[1].text = safe_unicode(co)

    doc.add_paragraph("")

    # ---------------- RBT TABLE ----------------
    doc.add_paragraph("RBT Classification")

    rbt = doc.add_table(rows=2, cols=7)
    rbt.style = "Table Grid"

    rbt.rows[0].cells[0].text = "Level"
    for idx, lvl in enumerate(["L1", "L2", "L3", "L4", "L5", "L6"], start=1):
        rbt.rows[0].cells[idx].text = lvl

    rbt.rows[1].cells[0].text = "Description"
    for idx, nm in enumerate(["Remembering","Understanding","Applying","Analyzing","Evaluating","Creating"], start=1):
        rbt.rows[1].cells[idx].text = nm

    doc.add_paragraph("")

    # ---------------- KC TABLE ----------------
    doc.add_paragraph("Knowledge Category")

    kc = doc.add_table(rows=2, cols=5)
    kc.style = "Table Grid"

    kc.rows[0].cells[0].text = "KC Code"
    for idx, code in enumerate(["F","C","P","M"], start=1):
        kc.rows[0].cells[idx].text = code

    kc.rows[1].cells[0].text = "KC Name"
    for idx, name in enumerate(["Factual","Conceptual","Procedural","Meta-Cognitive"], start=1):
        kc.rows[1].cells[idx].text = name

    # ---------------- RETURN ----------------
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()



# ============================================================
# STREAMLIT UI
# ============================================================
st.set_page_config(page_title="SmartTest Architect — Final", layout="wide")
st.title("📘 SmartTest Architect — MST Generator (Final)")

with st.sidebar:
    st.header("Teacher Details")
    university = st.text_input("University", "Maharaja Ranjit Singh Punjab Technical University")
    college = st.text_input("College", "GZS College of Engineering & Technology")
    department = st.text_input("Department", "Computer Science & Engineering")
    academic_year = st.text_input("Academic Year", "2025–2026")
    session = st.text_input("Session", "AUG–DEC 2025")
    program = st.text_input("Program", "B.Tech CSE AIML")
    semester = st.text_input("Semester", "1st")
    course_code = st.text_input("Course Code", "CSE101")
    course_title = st.text_input("Course Title", "Introduction to Computing")
    max_marks = st.text_input("Max Marks", "24")
    time_duration = st.text_input("Time Duration", "1.5 Hours")
    date_of_MST = st.text_input("Date of MST", "12-11-2025")
    roll_section = st.text_input("Roll No & Section", "")

    st.subheader("Course Outcomes")
    co_input = st.text_area("One CO per line", "1. Understand basics\n2. Apply logic")


syllabus_input = st.text_area("Paste Full Syllabus", "UNIT – I\nBasics...\n\nUNIT – II\nMore topics...", height=260)

units = parse_units_from_syllabus(syllabus_input)
co_lines = parse_co_list(co_input)

st.info(f"Detected {len(units)} units.")

# ----------------- DIFFICULTY -----------------
st.subheader("Difficulty Distribution")
colA, colB, colC = st.columns(3)
with colA:
    st.write("### Section A (4 questions)")
    A_easy = st.number_input("Easy", 0, 4, 1)
    A_med = st.number_input("Medium", 0, 4, 2)
    A_hard = st.number_input("Hard", 0, 4, 1)
with colB:
    st.write("### Section B (2 questions)")
    B_easy = st.number_input("Easy ", 0, 2, 0)
    B_med = st.number_input("Medium ", 0, 2, 1)
    B_hard = st.number_input("Hard ", 0, 2, 1)
with colC:
    st.write("### Section C (1 long)")
    C_easy = st.number_input("Easy  ", 0, 1, 0)
    C_med = st.number_input("Medium  ", 0, 1, 0)
    C_hard = st.number_input("Hard  ", 0, 1, 1)

# ----------------- QUESTION TYPE DISTRIBUTION (NEW) -----------------
st.subheader("Question Type Distribution (per section)")
st.write("**Section A (total must be 4)**")
A_logical = st.number_input("Section A - Logical", 0, 4, 0, key="A_logical")
A_syntax = st.number_input("Section A - Syntax", 0, 4, 0, key="A_syntax")
A_semantic = st.number_input("Section A - Semantic", 0, 4, 0, key="A_semantic")
A_theory = st.number_input("Section A - Theory", 0, 4, 0, key="A_theory")

st.write("**Section B (total must be 2)**")
B_logical = st.number_input("Section B - Logical", 0, 2, 0, key="B_logical")
B_syntax = st.number_input("Section B - Syntax", 0, 2, 0, key="B_syntax")
B_semantic = st.number_input("Section B - Semantic", 0, 2, 0, key="B_semantic")
B_theory = st.number_input("Section B - Theory", 0, 2, 0, key="B_theory")

st.write("**Section C (total must be 1)**")
C_logical = st.number_input("Section C - Logical", 0, 1, 0, key="C_logical")
C_syntax = st.number_input("Section C - Syntax", 0, 1, 0, key="C_syntax")
C_semantic = st.number_input("Section C - Semantic", 0, 1, 0, key="C_semantic")
C_theory = st.number_input("Section C - Theory", 0, 1, 0, key="C_theory")

# ----------------- NEW SECTION-WISE PER-UNIT -----------------
st.subheader("Section-wise Unit Distribution")
sectionA_unit = []; sectionB_unit = []; sectionC_unit = []

st.write("### Section A (4×2)")
for i, u in enumerate(units):
    sectionA_unit.append(st.number_input(f"{u['title']} in Section A", 0, 4, 0, key=f"A{i}"))

st.write("### Section B (2×4)")
for i, u in enumerate(units):
    sectionB_unit.append(st.number_input(f"{u['title']} in Section B", 0, 2, 0, key=f"B{i}"))

st.write("### Section C (1×8)")
for i, u in enumerate(units):
    sectionC_unit.append(st.number_input(f"{u['title']} in Section C", 0, 1, 0, key=f"C{i}"))

# ----------------- VALIDATION -----------------
valid = True
if sum(sectionA_unit) != 4:
    st.error("Section A must total 4 questions (unit distribution).")
    valid = False
if sum(sectionB_unit) != 2:
    st.error("Section B must total 2 questions (unit distribution).")
    valid = False
if sum(sectionC_unit) != 1:
    st.error("Section C must total 1 question (unit distribution).")
    valid = False

# validate question-type sums
if (A_logical + A_syntax + A_semantic + A_theory) != 4:
    st.error("Section A: Question Type counts must add to 4.")
    valid = False
if (B_logical + B_syntax + B_semantic + B_theory) != 2:
    st.error("Section B: Question Type counts must add to 2.")
    valid = False
if (C_logical + C_syntax + C_semantic + C_theory) != 1:
    st.error("Section C: Question Type counts must add to 1.")
    valid = False

# ============================================================
# GENERATE MST (button scope - ALL outputs built here)
# ============================================================
if st.button("⚡ Generate MST"):
    if not valid:
        st.error("Fix Section-wise distribution or Question Type counts first.")
    else:
        st.session_state["USED_Q"] = set()
        st.session_state["USED_TOPICS"] = set()

        sections = {"A": [], "B": [], "C": []}
        ans_pairs = []

        diffA = ["Easy"] * A_easy + ["Medium"] * A_med + ["Hard"] * A_hard
        diffB = ["Easy"] * B_easy + ["Medium"] * B_med + ["Hard"] * B_hard
        diffC = ["Easy"] * C_easy + ["Medium"] * C_med + ["Hard"] * C_hard
        random.shuffle(diffA); random.shuffle(diffB)
        if not diffC:
            diffC = ["Hard"]

        seqA = (["Logical"] * A_logical + ["Syntax"] * A_syntax + ["Semantic"] * A_semantic + ["Theory"] * A_theory)
        seqB = (["Logical"] * B_logical + ["Syntax"] * B_syntax + ["Semantic"] * B_semantic + ["Theory"] * B_theory)
        seqC = (["Logical"] * C_logical + ["Syntax"] * C_syntax + ["Semantic"] * C_semantic + ["Theory"] * C_theory)

        random.shuffle(seqA)
        random.shuffle(seqB)
        random.shuffle(seqC)

        ptrA = 0
        ptrB = 0
        ptrC = 0

        # SECTION A
        for idx, count in enumerate(sectionA_unit):
            for _ in range(count):
                if diffA:
                    diff = diffA[len(sections["A"]) % len(diffA)]
                else:
                    diff = "Medium"
                q_type = seqA.pop(0) if seqA else 'Generic'
                qdict = generate_questions(units[idx]["title"], units[idx]["content"], 1, kind="short", difficulty=diff, co_lines=co_lines, question_type=q_type)[0]
                sections["A"].append({"q": qdict["q"], "marks": 2, "tag": qdict["tag"]})
                ans_pairs.append((qdict["q"], generate_answer(qdict["q"], "short")))

        # SECTION B
        for idx, count in enumerate(sectionB_unit):
            for _ in range(count):
                if diffB:
                    diff = diffB[len(sections["B"]) % len(diffB)]
                else:
                    diff = "Medium"
                q_type = seqB.pop(0) if seqB else 'Generic'
                qdict = generate_questions(units[idx]["title"], units[idx]["content"], 1, kind="medium", difficulty=diff, co_lines=co_lines, question_type=q_type)[0]
                sections["B"].append({"q": qdict["q"], "marks": 4, "tag": qdict["tag"]})
                ans_pairs.append((qdict["q"], generate_answer(qdict["q"], "medium")))

        # SECTION C
        for idx, count in enumerate(sectionC_unit):
            if count == 1:
                diff = diffC[0] if diffC else "Hard"
                q_type = seqC.pop(0) if seqC else 'Generic'
                qdict = generate_questions(units[idx]["title"], units[idx]["content"], 1, kind="long", difficulty=diff, co_lines=co_lines, question_type=q_type)[0]
                sections["C"].append({"q": qdict["q"], "marks": 8, "tag": qdict["tag"]})
                ans_pairs.append((qdict["q"], generate_answer(qdict["q"], "long")))

        header = {
            "university": university,
            "college": college,
            "department": department,
            "academic_year": academic_year,
            "session": session,
            "program": program,
            "semester": semester,
            "course_code": course_code,
            "course_title": course_title,
            "mid_mst": "3rd",
            "coordinator": "Dr. Harish Kumar Garg",
            "max_marks": max_marks,
            "time_duration": time_duration,
            "date_of_MST": date_of_MST,
            "roll_number": roll_section,
            "co_lines": co_lines
        }

        # build pdfs & docx (call inside button block so 'sections' and 'header' exist)
        pdf_bytes = create_mst_pdf(sections, header)
        mst_docx = create_mst_docx(sections, header)

        st.session_state["mst_pdf"] = pdf_bytes
        st.session_state["mst_docx"] = mst_docx



        st.success("MST Generated Successfully 🎉")
        st.caption("⭐ Built by Kunal — SmartTest Architect (Final Clean Version)")

# ----------------- ALWAYS SHOW DOWNLOADS UNTIL REFRESH -----------------
if "mst_pdf" in st.session_state:
    st.subheader("📥 Download Your Generated MST")
    st.download_button("📄 MST Question Paper (PDF)", st.session_state["mst_pdf"], "MST_Questions.pdf")
    st.download_button("📝 MST Question Paper (DOCX)", st.session_state["mst_docx"], "MST_Questions.docx")

# footer caption (non-critical)
# st.caption("⭐ Built by Kunal — SmartTest Architect (Final Clean Version)")
